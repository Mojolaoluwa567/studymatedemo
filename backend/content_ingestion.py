"""
Content ingestion for all supported upload/input types.

Every extractor returns (text, page_count) - "page_count" is a loose
concept reused across types (PDF page count, 1 for plain text/web page,
estimated page-equivalent for Word) purely so the existing Document model
and UI ("X pages") don't need a source-type-specific field. The important
thing downstream (quiz generation, study aids) only ever sees `text`.

PDF extraction is unchanged from the original pdf_utils.py and kept in
its own function for clarity; this module is the single place new source
types get added going forward.
"""

import re
import tempfile
import os as _os

import pdfplumber
import docx
import requests
from bs4 import BeautifulSoup

from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api._errors import CouldNotRetrieveTranscript

# Cap how much text we send to the LLM per quiz generation request.
# Roughly ~4 chars per token, so this keeps us well within context limits
# while still giving the model plenty of material to draw from.
MAX_CHARS = 18000

# Words per "page" used to estimate a page-equivalent count for source
# types that don't have a native page concept (Word docs, web pages).
WORDS_PER_PAGE = 500


def extract_text_from_pdf(file_stream):
    """
    Extract text from a PDF file-like object.
    Returns (text, page_count).
    """
    text_parts = []
    page_count = 0

    with pdfplumber.open(file_stream) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    full_text = "\n\n".join(text_parts).strip()

    if not full_text:
        raise ValueError(
            "No extractable text found in this PDF. "
            "Scanned/image-only PDFs aren't supported yet."
        )

    return full_text, page_count


def extract_text_from_docx(file_stream):
    """
    Extract text from a Word (.docx) file-like object.
    Returns (text, estimated_page_count).

    Note: only .docx is supported (the modern XML-based format), not the
    legacy .doc binary format - python-docx can't read .doc files.
    """
    document = docx.Document(file_stream)

    parts = [p.text for p in document.paragraphs if p.text.strip()]

    # Also pull text out of any tables, since lecture notes/handouts often
    # use tables for definitions, comparisons, etc.
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    full_text = "\n\n".join(parts).strip()

    if not full_text:
        raise ValueError(
            "No extractable text found in this Word document. "
            "Make sure it isn't empty or image-only."
        )

    estimated_pages = max(1, len(full_text.split()) // WORDS_PER_PAGE)
    return full_text, estimated_pages


def extract_text_from_plain_text(raw_text):
    """
    Accepts already-decoded plain text (e.g. pasted notes, a .txt upload).
    Returns (text, estimated_page_count).
    """
    full_text = raw_text.strip()

    if not full_text:
        raise ValueError("No text was provided.")

    estimated_pages = max(1, len(full_text.split()) // WORDS_PER_PAGE)
    return full_text, estimated_pages


def extract_text_from_url(url):
    """
    Fetches a web page and extracts its main readable text.
    Returns (text, estimated_page_count).

    Deliberately simple: strips script/style/nav/footer/header tags and
    returns the remaining visible text. Good enough for articles, docs
    pages, and blog posts; won't handle JS-rendered single-page apps
    (no headless browser here) - that's a known limitation.
    """
    try:
        response = requests.get(
            url,
            timeout=10,
            headers={"User-Agent": "Mozilla/5.0 (StudyMate content fetcher)"},
        )
        response.raise_for_status()
    except requests.RequestException as e:
        raise ValueError(f"Could not fetch this URL: {e}")

    soup = BeautifulSoup(response.text, "html.parser")

    for tag in soup(["script", "style", "nav", "footer", "header", "noscript", "aside"]):
        tag.decompose()

    full_text = soup.get_text(separator="\n")
    # Collapse excessive blank lines left behind by stripped tags.
    lines = [line.strip() for line in full_text.splitlines()]
    full_text = "\n".join(line for line in lines if line).strip()

    if not full_text or len(full_text) < 50:
        raise ValueError(
            "Could not extract meaningful text from this page. "
            "It may require JavaScript to render, or have no readable content."
        )

    estimated_pages = max(1, len(full_text.split()) // WORDS_PER_PAGE)
    return full_text, estimated_pages


def truncate_for_prompt(text, max_chars=MAX_CHARS):
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\n\n[...content truncated...]"


# ---------------------------------------------------------------------------
# YouTube
# ---------------------------------------------------------------------------

_YOUTUBE_ID_PATTERNS = [
    r"(?:youtube\.com\/watch\?v=|youtube\.com\/shorts\/|youtu\.be\/|youtube\.com\/embed\/)([A-Za-z0-9_-]{11})",
]


def extract_youtube_video_id(url):
    """Pulls the 11-character video ID out of any common YouTube URL shape
    (watch, youtu.be short link, /shorts/, /embed/). Returns None if the
    URL doesn't look like a YouTube link at all."""
    for pattern in _YOUTUBE_ID_PATTERNS:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None


def extract_text_from_youtube(url, transcribe_audio_fn=None):
    """
    Extracts text from a YouTube video, in two stages:
      1. Try the video's captions (auto-generated or creator-uploaded) via
         youtube_transcript_api - fast, free, no audio processing.
      2. If no captions exist at all, fall back to downloading the audio
         with yt-dlp and transcribing it via the function passed in as
         `transcribe_audio_fn` (kept as an injected dependency so this
         function doesn't import quiz_generator/Gemini directly and stays
         independently testable).

    Returns (text, estimated_page_count).
    """
    video_id = extract_youtube_video_id(url)
    if not video_id:
        raise ValueError(
            "That doesn't look like a YouTube URL. Expected a youtube.com "
            "or youtu.be link."
        )

    # --- Stage 1: captions ---
    try:
        transcript = YouTubeTranscriptApi().fetch(video_id)
        full_text = " ".join(snippet.text for snippet in transcript).strip()
        if full_text:
            estimated_pages = max(1, len(full_text.split()) // WORDS_PER_PAGE)
            return full_text, estimated_pages
    except CouldNotRetrieveTranscript:
        pass  # fall through to audio fallback below
    except Exception:
        pass  # any other transcript-fetch hiccup also falls through

    # --- Stage 2: no captions - download audio and transcribe it ---
    if transcribe_audio_fn is None:
        raise ValueError(
            "This video has no captions available, and audio transcription "
            "isn't configured."
        )

    audio_bytes, mime_type = _download_youtube_audio(url)
    full_text = transcribe_audio_fn(audio_bytes, mime_type)

    if not full_text or not full_text.strip():
        raise ValueError(
            "Could not get a transcript for this video, even after "
            "attempting audio transcription."
        )

    full_text = full_text.strip()
    estimated_pages = max(1, len(full_text.split()) // WORDS_PER_PAGE)
    return full_text, estimated_pages


def _download_youtube_audio(url):
    """
    Downloads just the audio track of a YouTube video using yt-dlp, to a
    temp file, and returns (audio_bytes, mime_type). Used only as the
    fallback when no captions exist - this is meaningfully slower and
    heavier than the captions path, by design only hit when necessary.
    """
    import yt_dlp

    with tempfile.TemporaryDirectory() as tmp_dir:
        output_template = _os.path.join(tmp_dir, "audio.%(ext)s")
        ydl_opts = {
            "format": "worstaudio/worst",  # smallest file that still has audio - we only need speech, not quality
            "outtmpl": output_template,
            "quiet": True,
            "no_warnings": True,
            "noplaylist": True,
        }

        try:
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                ydl.download([url])
        except Exception as e:
            raise ValueError(f"Could not download audio for this video: {e}")

        downloaded_files = _os.listdir(tmp_dir)
        if not downloaded_files:
            raise ValueError("Audio download produced no file.")

        audio_path = _os.path.join(tmp_dir, downloaded_files[0])
        ext = _os.path.splitext(audio_path)[1].lstrip(".").lower()
        mime_type = AUDIO_MIME_TYPES.get(ext, "audio/mpeg")

        with open(audio_path, "rb") as f:
            audio_bytes = f.read()

    return audio_bytes, mime_type


# ---------------------------------------------------------------------------
# Audio file upload
# ---------------------------------------------------------------------------

AUDIO_MIME_TYPES = {
    "mp3": "audio/mpeg",
    "wav": "audio/wav",
    "m4a": "audio/mp4",
    "ogg": "audio/ogg",
    "webm": "audio/webm",
    "flac": "audio/flac",
    "aac": "audio/aac",
}

# Keep uploaded audio files reasonably small - long lecture recordings sent
# whole to Gemini cost more and take longer; this is a generous cap (good
# for ~1-2 hours of compressed speech) without being unbounded.
MAX_AUDIO_BYTES = 25 * 1024 * 1024  # 25MB


def extract_text_from_audio(file_bytes, filename, transcribe_audio_fn):
    """
    Transcribes an uploaded audio file via the injected transcribe_audio_fn
    (Gemini, in practice - see study_aids-adjacent transcribe_audio() in
    quiz_generator.py). Returns (text, estimated_page_count).
    """
    if len(file_bytes) > MAX_AUDIO_BYTES:
        raise ValueError(
            f"Audio file is too large ({len(file_bytes) // (1024*1024)}MB). "
            f"Max size is {MAX_AUDIO_BYTES // (1024*1024)}MB."
        )

    ext = _os.path.splitext(filename.lower())[1].lstrip(".")
    mime_type = AUDIO_MIME_TYPES.get(ext)
    if not mime_type:
        raise ValueError(
            f"Unsupported audio format '.{ext}'. Supported: "
            + ", ".join(f".{e}" for e in AUDIO_MIME_TYPES)
        )

    full_text = transcribe_audio_fn(file_bytes, mime_type)

    if not full_text or not full_text.strip():
        raise ValueError("Transcription returned no text for this audio file.")

    full_text = full_text.strip()
    estimated_pages = max(1, len(full_text.split()) // WORDS_PER_PAGE)
    return full_text, estimated_pages
